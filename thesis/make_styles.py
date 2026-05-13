"""Generate reference.docx with AITU thesis styles + title_page.docx with proper
AITU title page. Both are used by build.py.

Styling target (from AITU LaTeX template `memoirthesis.tex`):
- A4, 12pt Times New Roman, 1.5 line spacing
- Margins: top 40mm, bottom 30mm, left 30mm, right 20mm
- Page numbers: bottom centre
- Heading sizes scaled appropriately for Word equivalent of LaTeX memoir class
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

HERE = Path(__file__).parent
LOGO = HERE / "AITU.png"


def _set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def _set_page(section):
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)


def _add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _set_font(run, size=11)
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instr)
    run._element.append(fldChar2)


def _configure_style(doc, name, *, size, bold=False, italic=False,
                     align=None, space_before=0, space_after=6,
                     line_spacing=1.5):
    style = doc.styles[name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    # Force font also at rPr level for older Word compatibility
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Times New Roman")
    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def make_reference_docx(out: Path) -> None:
    """A pandoc reference document. pandoc reads styles from here and applies
    them when converting Markdown → docx."""
    doc = Document()
    _set_page(doc.sections[0])
    _add_page_number_footer(doc.sections[0])

    _configure_style(doc, "Normal", size=12,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                     space_after=6)
    _configure_style(doc, "Heading 1", size=16, bold=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=18, space_after=12, line_spacing=1.15)
    _configure_style(doc, "Heading 2", size=14, bold=True,
                     space_before=14, space_after=8, line_spacing=1.15)
    _configure_style(doc, "Heading 3", size=13, bold=True,
                     space_before=10, space_after=6, line_spacing=1.15)
    _configure_style(doc, "Heading 4", size=12, bold=True, italic=True,
                     space_before=8, space_after=4, line_spacing=1.15)
    _configure_style(doc, "Title", size=20, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=0, space_after=24, line_spacing=1.15)

    # Ensure a Hyperlink style exists (pandoc references it)
    try:
        _configure_style(doc, "Hyperlink", size=12,
                         line_spacing=1.5, space_after=0)
    except KeyError:
        pass

    doc.save(str(out))


def _hr(p):
    """Insert a horizontal rule under a paragraph (used for title page rules)."""
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _centered_para(doc, text=None, *, size=12, bold=False, italic=False,
                   space_after=0, line_spacing=1.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text is not None:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
    return p


def make_title_page(out: Path) -> None:
    """Build a standalone title.docx mimicking the AITU LaTeX titlingpage.
    Engineered to fit exactly 1 A4 page."""
    doc = Document()
    sec = doc.sections[0]
    _set_page(sec)
    # Tighter top margin for title page so everything fits
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)

    # Header text — institution
    _centered_para(doc, "LIMITED LIABILITY PARTNERSHIP", size=11, bold=True,
                   space_after=2)
    _centered_para(doc, "ASTANA IT UNIVERSITY", size=11, bold=True,
                   space_after=12)

    # Double horizontal rule
    _hr(_centered_para(doc, space_after=2))
    _hr(_centered_para(doc, space_after=8))

    # Document type
    _centered_para(doc, "DIPLOMA PROJECT", size=22, bold=True, space_after=18)

    # Title
    _centered_para(doc, "Topic:", size=12, bold=True, space_after=4)
    _centered_para(doc, "Tree Detection for Astana —", size=18, bold=True,
                   space_after=2, line_spacing=1.2)
    _centered_para(doc, "Deep Learning for Urban Green Space Mapping",
                   size=18, bold=True, space_after=14, line_spacing=1.2)

    # Double horizontal rule
    _hr(_centered_para(doc, space_after=2))
    _hr(_centered_para(doc, space_after=14))

    # AITU logo (small)
    if LOGO.exists():
        p = _centered_para(doc, space_after=10)
        run = p.add_run()
        run.add_picture(str(LOGO), width=Cm(3))

    # "By" label
    _centered_para(doc, "By", size=12, italic=True, space_after=4)

    # Authors
    for full in ("Totin Anuar", "Aidarkhanov Rasul", "Sharipov Berik"):
        _centered_para(doc, full, size=13, bold=True, space_after=2)

    _centered_para(doc, space_after=12)  # spacer

    # Program + supervisor
    _centered_para(doc, "Educational Program: 6B06103 — Information Technologies",
                   size=11, space_after=4)
    _centered_para(doc, "Scientific Supervisor: Syndar Satbayev",
                   size=11, space_after=4)
    _centered_para(doc, "School of Information Technologies, Astana IT University",
                   size=11, italic=True, space_after=10)

    # Place + year
    _centered_para(doc, "ASTANA, 2026", size=12, bold=True, space_after=0)

    # Page break so content starts on next page
    doc.add_page_break()

    doc.save(str(out))


def main():
    ref = HERE / "reference.docx"
    title = HERE / "title_page.docx"
    make_reference_docx(ref)
    print(f"Wrote {ref} ({ref.stat().st_size:,} bytes)")
    make_title_page(title)
    print(f"Wrote {title} ({title.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
